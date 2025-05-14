import tkinter as tk
from tkinter import ttk, messagebox
from tkinter import filedialog
from tkinter import PhotoImage
import openpyxl
import re
from docx import Document

DB_PATH = "C:\DataBase\DB2.xlsx"

workbook = openpyxl.load_workbook(DB_PATH)
users_sheet = workbook["Users"]
equipment_sheet = workbook["Equipment"]

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Аренда фототехники")
        self.geometry("655x600")
        self.resizable(False, False)
        self.configure(bg="#CB7562") 

        self.frames = {}

        # Инициализируем все страницы
        for F in (HomePage, RegistrationPage, LoginPage, RentalPage, SuccessPage):
            page_name = F.__name__
            frame = F(parent=self, controller=self)
            self.frames[page_name] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        # Показываем начальную страницу
        self.show_frame("HomePage")

    def show_frame(self, page_name):
        """Отображает страницу по имени."""
        frame = self.frames[page_name]
        frame.tkraise()

        # Очистка полей на странице входа, если возвращаемся на нее
        if page_name == "LoginPage":
            frame.clear_fields()  # Очищаем поля входа при возвращении на страницу

class CenteredFrame(tk.Frame):
    """Базовый класс для центрирования содержимого."""
    def __init__(self, parent, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)
    
class HomePage(CenteredFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        frame = tk.Frame(self)
        frame.pack(fill="both", expand=True)

        image_path = r"C:\DataBase\Component 1.png"  # Путь к изображению
        img = tk.PhotoImage(file=image_path)

        # Уменьшаем изображение в 2 раза
        resized_img = img.subsample(3, 3)  # Значения 2, 2 - коэффициенты уменьшения по горизонтали и вертикали

        header = tk.Label(
            frame,
            image=resized_img,
            text="PHotoNN",
            font=("Verdana", 25, "bold"),
            bg="#CB7562",
            fg="#FFFFFF",
            compound="left",
            padx=20,
            height=100
        )
        header.image = resized_img  # Сохраняем ссылку на уменьшенное изображение
        header.pack(side="top", fill="x", pady=10)

        tk.Label(frame, text="Добро пожаловать!", font=("Verdana", 25, "bold")).pack(pady=50)

        tk.Button(frame, bg="#CB7562", fg="#FFFFFF", text="Регистрация", font=("Verdana", 18, "bold"),
                  command=lambda: controller.show_frame("RegistrationPage")).pack(pady=10)

        tk.Button(frame, bg="#CB7562", fg="#FFFFFF", text="Вход", font=("Verdana", 18, "bold"),
                  command=lambda: controller.show_frame("LoginPage")).pack(pady=10)

        tk.Button(frame, bg="#808080", fg="#FFFFFF", text="Выйти из приложения", font=("Verdana", 12, "bold"),
                  command=self.confirm_exit).pack(pady=30)
    
    def confirm_exit(self):
        if tk.messagebox.askyesno("Выход", "Вы действительно хотите выйти из приложения?"):
            self.controller.quit()

class LoginPage(CenteredFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller  # Контроллер

        frame = tk.Frame(self)
        frame.grid(row=0, column=0, padx=20, pady=20)

        tk.Label(frame, text="Вход", font=("Verdana", 25, "bold")).pack(pady=40)

        tk.Label(frame, text="Email:", font=("Verdana", 12)).pack(anchor="w")
        self.email_entry = tk.Entry(frame, font=("Verdana", 12), width=40)
        self.email_entry.pack(pady=5)

        tk.Label(frame, text="Пароль:", font=("Verdana", 12)).pack(anchor="w")
        self.password_entry = tk.Entry(frame, show="*", font=("Verdana", 12), width=40)
        self.password_entry.pack(pady=5)

        self.login_button = tk.Button(frame, bg="#CB7562", fg="#FFFFFF", text="    Войти    ", font=("Verdana", 20, "bold"), command=self.login_user)
        self.login_button.pack(pady=10)

        tk.Button(frame, bg="#CB7562", fg="#FFFFFF", text="На главную", font=("Verdana", 12, "bold"),
                  command=lambda: controller.show_frame("HomePage")).pack(pady=10)

        # Метки для отображения ошибок
        self.error_label = tk.Label(frame, text="", font=("Verdana", 10), fg="red")
        self.error_label.pack()

    def clear_fields(self):
        """Очищает поля email и пароль при возврате на страницу входа."""
        self.email_entry.delete(0, tk.END)
        self.password_entry.delete(0, tk.END)
        self.error_label.config(text="")  # Очищаем текст ошибки

    def login_user(self):
        email = self.email_entry.get().strip()
        password = self.password_entry.get().strip()

        # Проверка на пустые поля
        if not email or not password:
            self.error_label.config(text="Пожалуйста, заполните все поля!")
            return

        # Проверка входа в базе данных
        user_found = False
        for row in users_sheet.iter_rows(min_row=2, values_only=True):
            if row[2] == email and row[3] == password:
                user_found = True
                break

        if user_found:
            self.error_label.config(text="")  # Очищаем текст ошибки
            messagebox.showinfo("Вход", "Успешный вход!")
            self.controller.show_frame("RentalPage")
        else:
            self.error_label.config(text="Неверный email или пароль.")  # Отображаем ошибку, если вход не успешен
    
    custom_font = ("Verdana", 20)

class RegistrationPage(CenteredFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        frame = tk.Frame(self)
        frame.grid(row=0, column=0, padx=20, pady=20)

        tk.Label(frame, text="Регистрация", font=("Verdana", 25, "bold")).pack(pady=40)

        tk.Label(frame, text="Имя пользователя:", font=("Verdana", 12)).pack(anchor="w")
        self.name_entry = tk.Entry(frame, font=("Verdana", 12), width=40)
        self.name_entry.pack(pady=20)

        tk.Label(frame, text="Email:", font=("Verdana", 12)).pack(anchor="w")
        self.email_entry = tk.Entry(frame, font=("Verdana", 12), width=40)
        self.email_entry.pack(pady=5)
        self.email_entry.bind("<KeyRelease>", self.validate_email_realtime)  # Реальная проверка email
        self.email_error_label = tk.Label(frame, text="", font=("Verdana", 10), fg="red")
        self.email_error_label.pack(anchor="w")

        tk.Label(frame, text="Пароль:", font=("Verdana", 12)).pack(anchor="w")
        self.password_entry = tk.Entry(frame, show="*", font=("Verdana", 12), width=40)
        self.password_entry.pack(pady=5)
        self.password_entry.bind("<KeyRelease>", self.validate_password_realtime)  # Реальная проверка пароля
        self.password_error_label = tk.Label(frame, text="", font=("Verdana", 10), fg="red")
        self.password_error_label.pack(anchor="w")

        tk.Button(frame, bg="#CB7562", fg="#FFFFFF", text="  Зарегистрироваться  ", font=("Verdana", 20, "bold"), command=self.register_user).pack(pady=10)

        tk.Button(frame, bg="#CB7562", fg="#FFFFFF", text="На главную", font=("Verdana", 12, "bold"),
                  command=lambda: controller.show_frame("HomePage")).pack(pady=10)

    def validate_email(self, email):
        """Проверяет, является ли email корректным."""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email))

    def validate_email_realtime(self, event=None):
        """Проверяет email в реальном времени."""
        email = self.email_entry.get().strip()
        if not self.validate_email(email):
            self.email_error_label.config(text="Некорректный email.")
        else:
            self.email_error_label.config(text="")

    def validate_password(self, password):
        """Проверяет, соответствует ли пароль минимальным требованиям."""
        if len(password) < 8:
            return "Пароль должен быть длиной не менее 8 символов."
        if not any(char.isdigit() for char in password):
            return "Пароль должен содержать хотя бы одну цифру."
        if not any(char.isalpha() for char in password):
            return "Пароль должен содержать хотя бы одну букву."
        return None

    def validate_password_realtime(self, event=None):
        """Проверяет пароль в реальном времени."""
        password = self.password_entry.get().strip()
        error = self.validate_password(password)
        if error:
            self.password_error_label.config(text=error)
        else:
            self.password_error_label.config(text="")

    def register_user(self):
        name = self.name_entry.get().strip()
        email = self.email_entry.get().strip()
        password = self.password_entry.get().strip()

        # Проверка заполнения всех полей
        if not name or not email or not password:
            messagebox.showwarning("Ошибка", "Заполните все поля")
            return

        # Валидация email
        if not self.validate_email(email):
            messagebox.showwarning("Ошибка", "Введите корректный email")
            return

        # Валидация пароля
        password_error = self.validate_password(password)
        if password_error:
            messagebox.showwarning("Ошибка", password_error)
            return

        # Сохранение данных пользователя
        new_row = [users_sheet.max_row, name, email, password]
        users_sheet.append(new_row)
        workbook.save(DB_PATH)

        messagebox.showinfo("Регистрация", "Успешная регистрация!")
        self.controller.show_frame("RentalPage")


class RentalPage(CenteredFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        frame = tk.Frame(self)
        frame.grid(row=0, column=0, padx=20, pady=20)

        tk.Label(frame, text="Аренда фототехники", font=("Verdana", 25, "bold")).pack(pady=10)

        # Поле для ввода поиска
        search_frame = tk.Frame(frame)
        search_frame.pack(fill=tk.X, pady=5)
        tk.Label(search_frame, text="Поиск:", font=("Verdana", 12)).pack(side="left", padx=5)
        self.search_entry = tk.Entry(search_frame, font=("Verdana", 12), width=25)
        self.search_entry.pack(side="left", padx=5)
        self.search_entry.bind("<KeyRelease>", self.filter_data)  # Фильтрация при вводе текста

        columns = ("Аппаратура", "Цена за день")
        self.tree = ttk.Treeview(frame, columns=columns, show="headings", height=10)
        for col in columns:
            self.tree.heading(col, text=col, anchor="w")
            self.tree.column(col, width=200 if col == "Аппаратура" else 100, anchor="w")

        self.tree.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        self.equipment_data = []  # Список для хранения всех данных об оборудовании
        self.load_data()

        tk.Label(frame, text="Количество дней:", font=("Verdana", 16)).pack(pady=5)
        self.days_entry = tk.Entry(frame, font=("Verdana", 12), width=10)
        self.days_entry.pack()

        self.total_label = tk.Label(frame, text="Итоговая сумма: 0 руб.", font=("Verdana", 16))
        self.total_label.pack(pady=10)

        button_frame = tk.Frame(frame)
        button_frame.pack(pady=10)

        self.calculate_button = tk.Button(button_frame, bg="#CB7562", fg="#FFFFFF", text="Рассчитать", font=("Verdana", 10, "bold"), command=self.calculate_cost)
        self.calculate_button.pack(side="left", padx=5)

        self.rent_button = tk.Button(button_frame, bg="#CB7562", fg="#FFFFFF", text="Арендовать", state="disabled", font=("Verdana", 10, "bold"), command=self.confirm_rent)
        self.rent_button.pack(side="left", padx=5)

        tk.Button(button_frame, text="Выйти из аккаунта", bg="#CB6262", fg="#FFFFFF", font=("Verdana", 10, "bold"), command=self.confirm_switch_account).pack(side="left", padx=5)

    def load_data(self):
        """Загружает данные оборудования из листа и сохраняет их для фильтрации."""
        self.equipment_data = []  # Очищаем старые данные
        for row in equipment_sheet.iter_rows(min_row=2, values_only=True):
            name, price = row[1], row[2]
            name = name.strip() if name else "—"
            price = price if price else 0
            self.equipment_data.append((name, price))  # Сохраняем данные
        self.filter_data()  # Обновляем отображение

    def filter_data(self, event=None):
        """Фильтрует данные на основе введенного поискового запроса."""
        query = self.search_entry.get().strip().lower().replace(" ", "")  # Убираем пробелы и приводим к нижнему регистру
        filtered_data = [
            item for item in self.equipment_data
            if query in item[0].lower().replace(" ", "")  # Сравниваем с учетом требований
        ]

        self.tree.delete(*self.tree.get_children())  # Очищаем дерево
        for name, price in filtered_data:
            self.tree.insert("", "end", values=(name, price))  # Отображаем только отфильтрованные данные

    def calculate_cost(self):
        days = self.days_entry.get()
        if not days.isdigit() or int(days) <= 0:
            messagebox.showwarning("Ошибка", "Введите корректное количество дней!")
            return

        days = int(days)
        total_cost = 0

        for item_id in self.tree.selection():
            item = self.tree.item(item_id)
            price_per_day = item["values"][1]
            total_cost += price_per_day * days

        self.total_label.config(text=f"Итоговая сумма: {total_cost} руб.")

        # Включаем кнопку "Арендовать"
        self.calculate_button_enabled = True
        self.rent_button.config(state="normal")

    def confirm_rent(self):
        if not self.calculate_button_enabled:
            messagebox.showwarning("Ошибка", "Сначала рассчитайте стоимость аренды.")
            return

        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("Ошибка", "Выберите оборудование для аренды")
            return

        confirmation = messagebox.askyesno("Аренда", "Подтвердите аренду")
        if confirmation:
            self.rent_items()

    def rent_items(self):
        days = self.days_entry.get()
        if not days.isdigit() or int(days) <= 0:
            messagebox.showwarning("Ошибка", "Введите корректное количество дней")
            return

        days = int(days)
        total_cost = 0

        for item_id in self.tree.selection():
            item = self.tree.item(item_id)
            price_per_day = item["values"][1]
            total_cost += price_per_day * days

        self.controller.show_frame("SuccessPage")
    
    def confirm_switch_account(self):
        """Подтверждение выхода из аккаунта с диалогом Да/Нет."""
        confirm_exit = messagebox.askyesno("Выход", "Вы уверены, что хотите выйти из аккаунта?")
        if confirm_exit:
            self.controller.show_frame("HomePage")  # Возвращаем на страницу входа


class SuccessPage(CenteredFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        frame = tk.Frame(self)
        frame.grid(row=0, column=0, padx=20, pady=20)

        tk.Label(frame, text="Аренда подтверждена!", font=("Verdana", 25, "bold")).pack(pady=10)
        tk.Label(frame, text="Наш специалист свяжется с вами в течение 15 минут для\nподтверждения аренды", font=("Verdana", 12)).pack(pady=10)
        tk.Label(frame, text="Наш адресс:", font=("Verdana", 16, "bold")).pack(pady=5)
        tk.Label(frame, text="г. Нижний новгород, ул. Пушкина, д. 10\nEmail: rental@company.com", font=("Verdana", 12)).pack()

        button_frame = tk.Frame(frame)
        button_frame.pack(pady=10)

        tk.Button(button_frame, bg="#CB7562", fg="#FFFFFF", text="Назад к аренде", font=("Verdana"), command=lambda: self.controller.show_frame("RentalPage")).pack(side="left", padx=5)
        tk.Button(button_frame, bg="#CB7562", fg="#FFFFFF", text="На главную", font=("Verdana"), command=lambda: self.controller.show_frame("HomePage")).pack(side="left", padx=5)
        tk.Button(button_frame, bg="#CB7562", fg="#FFFFFF", text="Экспорт в Word", font=("Verdana"), command=self.export_to_word).pack(side="left", padx=5)  # Кнопка экспорта
        tk.Button(button_frame, bg="#CB7562", fg="#FFFFFF", text="Выход", font=("Verdana"), command=self.exit_program).pack(side="left", padx=5)

    def exit_program(self):
        confirm_exit = messagebox.askyesno("Выход", "Вы уверены, что хотите выйти?")
        if confirm_exit:
            self.controller.quit()

    def export_to_word(self):
        """Экспортирует список арендованной аппаратуры в Word."""
        # Получаем данные из дерева на странице аренды
        rental_page = self.controller.frames["RentalPage"]
        selected_items = rental_page.tree.selection()
        if not selected_items:
            messagebox.showwarning("Ошибка", "Список аренды пуст.")
            return

        # Создаем документ Word
        document = Document()
        document.add_heading("Список арендованной аппаратуры", level=1)

        # Добавляем общую информацию
        document.add_paragraph(f"Итоговая сумма: {rental_page.total_label['text']}")
        document.add_paragraph(f"Количество дней: {rental_page.days_entry.get()}")

        # Таблица с выбранной аппаратурой
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Аппаратура"
        hdr_cells[1].text = "Цена за день"

        for item_id in selected_items:
            item = rental_page.tree.item(item_id)
            name, price = item["values"]
            row_cells = table.add_row().cells
            row_cells[0].text = str(name)
            row_cells[1].text = f"{price} руб."

        # Сохранение файла
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Документы Word", "*.docx")],
                title="Сохранить файл как",
                initialfile="Аренда_аппаратуры.docx"
            )
            if file_path:  # Проверка, выбрал ли пользователь путь
                document.save(file_path)
                messagebox.showinfo("Экспорт", f"Файл успешно сохранен как {file_path}")
            else:
                messagebox.showinfo("Экспорт", "Сохранение отменено")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {e}")

if __name__ == "__main__":
    app = App()
    app.mainloop()
